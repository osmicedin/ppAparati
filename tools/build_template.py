from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAME = "Arial"
TABLE_WIDTHS = [500, 650, 800, 1350, 850, 1100, 1100, 1450, 1400, 2000]


def set_run_font(run, size: float = 11, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_NAME)
    fonts.set(qn("w:hAnsi"), FONT_NAME)
    fonts.set(qn("w:eastAsia"), FONT_NAME)


def clear_paragraph_content(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_text(paragraph, text: str, size: float = 11, bold: bool = False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


def add_sdt_run(paragraph, tag: str, placeholder: str, size: float = 11, bold: bool = False) -> None:
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), tag)
    sdt_pr.append(alias)

    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    sdt_pr.append(tag_element)

    text_property = OxmlElement("w:text")
    sdt_pr.append(text_property)
    sdt.append(sdt_pr)

    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_NAME)
    fonts.set(qn("w:hAnsi"), FONT_NAME)
    fonts.set(qn("w:eastAsia"), FONT_NAME)
    rpr.append(fonts)
    if bold:
        rpr.append(OxmlElement("w:b"))
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    rpr.append(size_element)
    run.append(rpr)

    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = placeholder
    run.append(text)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def find_paragraph(document, predicate, description: str):
    matches = [paragraph for paragraph in document.paragraphs if predicate(paragraph.text.strip())]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one paragraph for {description}, found {len(matches)}")
    return matches[0]


def rebuild_slots(document) -> None:
    paragraph = find_paragraph(document, lambda text: text.startswith("Broj: 006-0430/26"), "report number")
    clear_paragraph_content(paragraph)
    add_text(paragraph, "Broj: ", size=12)
    add_sdt_run(paragraph, "ReportNumber", "BROJ_ZAPISNIKA", size=12)

    paragraph = find_paragraph(document, lambda text: text.startswith("Datum:") and "31.07.2026" in text, "conclusion date")
    clear_paragraph_content(paragraph)
    add_text(paragraph, "Datum: ", size=12, bold=True)
    add_sdt_run(paragraph, "ConclusionDate", "DATUM_ZAKLJUCIVANJA", size=12, bold=True)

    paragraph = find_paragraph(document, lambda text: text.startswith("ZA") and "HIFA" in text.upper(), "customer title")
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(paragraph, 'ZA "', size=13, bold=True)
    add_sdt_run(paragraph, "CustomerTitle", "NAZIV_KUPCA", size=13, bold=True)
    add_text(paragraph, '"', size=13, bold=True)

    paragraph = find_paragraph(document, lambda text: text.startswith("Doboj Jug, juli"), "location month year")
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_sdt_run(paragraph, "LocationMonthYear", "LOKACIJA_MJESEC_GODINA", size=12)

    paragraph = find_paragraph(document, lambda text: text.startswith("Naručilac:"), "orderer")
    clear_paragraph_content(paragraph)
    add_text(paragraph, "Naručilac:  ", size=11)
    add_sdt_run(paragraph, "CustomerOrderer", "NAZIV_NARUCIOCA", size=11, bold=True)

    paragraph = find_paragraph(document, lambda text: text.startswith("Na osnovu zahtjeva") and "09.07.2026" in text, "inspection period")
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text(
        paragraph,
        "Na osnovu zahtjeva, a shodno članu 39. Zakona o zaštiti od požara i vatrogastvu FBiH "
        "(Službene novine FBiH, broj 64/09), članu 33. Zakona o zaštiti požara Zeničko-dobojskog "
        "kantona (Službene novine ZE-DO kantona broj 5/11), te odredaba Pravilnika o izboru i "
        "održavanju aparata za gašenje početnog požara koji se mogu stavljati u promet sa garantnim "
        "rokom i rokom servisiranja (Službene novine FBiH broj 46/11), a u cilju unapređenja i "
        "sprovođenja zaštite od požara, u periodu ",
        size=11,
    )
    add_sdt_run(paragraph, "PeriodFrom", "PERIOD_OD", size=11, bold=True)
    add_text(paragraph, " do ", size=11, bold=True)
    add_sdt_run(paragraph, "PeriodTo", "PERIOD_DO", size=11, bold=True)
    add_text(
        paragraph,
        " izvršeno je periodično ispitivanje aparata za početno gašenje požara na osnovu čega "
        "izdajemo sljedeći:",
        size=11,
    )

    paragraph = find_paragraph(document, lambda text: text.startswith("u vlasništvu:") and "Hifa Oil" in text, "owner")
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(paragraph, 'u vlasništvu:  "', size=12)
    add_sdt_run(paragraph, "CustomerOwner", "NAZIV_VLASNIKA", size=12, bold=True)
    add_text(paragraph, '"', size=12)

    paragraph = find_paragraph(document, lambda text: text.startswith("Zaključeno sa rednim brojem"), "conclusion count")
    clear_paragraph_content(paragraph)
    add_text(paragraph, "Zaključeno sa rednim brojem ", size=11)
    add_sdt_run(paragraph, "ConclusionCount", "BROJ_STAVKI", size=11)
    add_text(paragraph, ".!", size=11)


def remove_next_service_notes(document) -> None:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip().upper().startswith("NAREDNO KONTROLNO ISPITIVANJE"):
            element = paragraph._element
            element.getparent().remove(element)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 7.5, center: bool = True) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    ppr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "180")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_table_geometry(table) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(sum(TABLE_WIDTHS)))

    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")
    tbl_pr.append(indent)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), "60")
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)

    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tbl_pr.append(borders)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in TABLE_WIDTHS:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table._tbl.findall(qn("w:tr")):
        grid_index = 0
        for cell in row.findall(qn("w:tc")):
            tc_pr = cell.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = OxmlElement("w:tcPr")
                cell.insert(0, tc_pr)
            grid_span = tc_pr.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            width_value = sum(TABLE_WIDTHS[grid_index : grid_index + span])
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(width_value))
            grid_index += span


def wrap_marker_row(table) -> None:
    marker_row = table.rows[2]._tr
    parent = marker_row.getparent()
    position = parent.index(marker_row)
    parent.remove(marker_row)

    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "ReportRows")
    sdt_pr.append(tag)
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "ReportRows")
    sdt_pr.append(alias)
    sdt.append(sdt_pr)

    content = OxmlElement("w:sdtContent")
    content.append(marker_row)
    sdt.append(content)
    parent.insert(position, sdt)


def replace_table(document) -> None:
    if len(document.tables) != 1:
        raise RuntimeError(f"Expected one source table, found {len(document.tables)}")

    source_table = document.tables[0]
    new_table = document.add_table(rows=3, cols=10)
    source_table._tbl.addprevious(new_table._tbl)
    source_table._tbl.getparent().remove(source_table._tbl)

    new_table.cell(0, 0).merge(new_table.cell(1, 0))
    new_table.cell(0, 1).merge(new_table.cell(0, 9))
    set_cell_text(new_table.cell(0, 0), "Redni\nbroj", bold=True)
    set_cell_text(new_table.cell(0, 1), "Identifikacioni podaci aparata", bold=True, size=8)

    headers = [
        "Tip",
        "Punjenje\nkg",
        "Serijski broj\naparata",
        "Godina\nproizvodnje",
        "Datum\nservisa",
        "Sljedeći\nservis",
        "Konstatacija\nispravnosti",
        "Vozilo",
        "Ispitivanje\nizvršio",
    ]
    for index, text in enumerate(headers, start=1):
        set_cell_text(new_table.cell(1, index), text, bold=True)

    for cell in new_table.rows[2].cells:
        set_cell_text(cell, "")
    set_cell_text(new_table.rows[2].cells[0], "{{REPORT_ROWS}}", size=7.5)

    set_repeat_table_header(new_table.rows[0])
    set_repeat_table_header(new_table.rows[1])
    set_table_geometry(new_table)
    wrap_marker_row(new_table)


def set_document_update_fields(document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build(reference: Path, output: Path) -> None:
    document = Document(reference)
    rebuild_slots(document)
    remove_next_service_notes(document)
    replace_table(document)
    set_document_update_fields(document)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = "Predložak zapisnika PP aparata"
    document.core_properties.subject = "Mjesečni izvještaj servisiranih PP aparata"
    document.core_properties.author = "ppEvidencija"
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("reference", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.reference.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
